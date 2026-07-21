from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import shutil

src = r"c:\Users\Jaedon Ong\Downloads\testplantemplatev2.docx"
out = r"c:\Users\Jaedon Ong\Downloads\HomeShare-Test-Plan.docx"
shutil.copy2(src, out)

doc = Document(out)
body = doc.element.body
for child in list(body):
    if child.tag != qn("w:sectPr"):
        body.remove(child)


def add_heading_like(text, bold=True, size=14):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p


def add_normal(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    return p


def add_case(num, intention, module, success, fail, survey_q):
    p = doc.add_paragraph()
    run = p.add_run(f"{num}.\tTested Intention: {intention}")
    run.bold = True
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.add_run(f"Module / area: {module}").font.size = Pt(11)

    p = doc.add_paragraph()
    p.add_run(f"Assessed by survey: {survey_q}").font.size = Pt(11)

    p = doc.add_paragraph()
    r = p.add_run("Success Criteria:")
    r.bold = True
    r.font.size = Pt(11)
    p = doc.add_paragraph(success)
    for run in p.runs:
        run.font.size = Pt(11)

    p = doc.add_paragraph()
    r = p.add_run("Fail Criteria:")
    r.bold = True
    r.font.size = Pt(11)
    p = doc.add_paragraph(fail)
    for run in p.runs:
        run.font.size = Pt(11)

    doc.add_paragraph()


add_heading_like("Test Planning Summary — HomeShare (IWP)", bold=True, size=16)
doc.add_paragraph()
add_normal("Product: HomeShare — local-network + cloud file sharing platform")
add_normal(
    "Test Duration (hours): ______ (suggested: 4–6 hours; survey filled right after testing)"
)
add_normal("Survey link:")
add_normal("[Add your Google Form link here — keep it live for the entire semester]")
doc.add_paragraph()

add_heading_like("How the survey assesses the test cases", bold=True, size=13)
add_normal(
    "Each survey question is numbered to match the test case in the same section "
    "(Unit question 1 → Unit case 1, and so on). Testers answer Yes / Partially / No. "
    "Mark the linked test case from the survey:"
)
add_normal("• Yes → Pass")
add_normal("• Partially → Pass with notes (or Fail if success criteria require full success)")
add_normal("• No → Fail")
add_normal(
    "Use open-comment questions only for evidence; Pass/Fail still comes from the Yes/Partial/No items."
)
doc.add_paragraph()

add_heading_like("Testing approach (Testing Pyramid)", bold=True, size=13)
add_normal("1) Unit — small, fast checks of individual logic/UI pieces")
add_normal("2) Integration — API, storage, and module interactions")
add_normal("3) End-to-end — full user workflows (cloud + local)")
doc.add_paragraph()

add_heading_like("List of Test Cases & Scenarios", bold=True, size=13)
doc.add_paragraph()

# -------- UNIT --------
add_heading_like("A. Unit Tests (bottom of pyramid)", bold=True, size=12)
add_normal(
    "Focus: Individual component testing; fast, isolated tests; core logic validation."
)
doc.add_paragraph()

unit = [
    (
        1,
        "Password rules reject weak passwords and accept valid ones",
        "Auth / passwordValidation",
        "Weak passwords are rejected with a clear message; a valid password is accepted on register/reset.",
        "Weak passwords accepted, or valid passwords blocked.",
        "Unit survey question 1",
    ),
    (
        2,
        "Access modes Private / Shared / Local Only behave as labelled",
        "Client / accessModes",
        "Private cannot create share links; Shared and Local Only can; labels match the chosen mode.",
        "Private allows sharing, or labels do not match the selected mode.",
        "Unit survey question 2",
    ),
    (
        3,
        "Cloud vs Local file badges and filters are understandable and correct",
        "Client / fileStorageScope + Library",
        "Badges show Cloud or Local; All / Cloud / Local filters change the visible set as expected.",
        "Everything looks Cloud, filters do nothing, or badges contradict the mode you uploaded in.",
        "Unit survey question 3",
    ),
    (
        4,
        "Header shows current mode (Cloud / Local) and Exit local only when Local is on",
        "Client / Header",
        "Badge matches the mode; Exit local appears only in Local mode.",
        "Wrong badge, or Exit local shown while on Cloud.",
        "Unit survey question 4",
    ),
    (
        5,
        "Help page explains Cloud vs Local and links to Local Network setup",
        "Client / Help",
        "Help is reachable from header/footer; explanations match the product; setup link/button works.",
        "Help missing, unclear, or cannot open Local Network setup from Help.",
        "Unit survey question 5",
    ),
    (
        6,
        "Idle timeout concept is clear (activity on HomeShare matters)",
        "Client / idleTimeout",
        "User understands they can be logged out after inactivity on the site (~15 minutes).",
        "No idea why logout happened, or never sees logout after long idle.",
        "Unit survey question 6",
    ),
]

for row in unit:
    add_case(*row)

# -------- INTEGRATION --------
add_heading_like("B. Integration Tests (middle of pyramid)", bold=True, size=12)
add_normal("Focus: System interaction; data flow; API and service integration.")
doc.add_paragraph()

integ = [
    (
        1,
        "Cloud login works without a local server",
        "Auth ↔ Cloud API",
        "On the live site (Cloud mode), login reaches the dashboard with no localhost / Detect required.",
        "Login blocked by local-server errors or connection refused to 127.0.0.1.",
        "Integration survey question 1",
    ),
    (
        2,
        "Detect on the host PC switches the site into Local mode",
        "apiDiscovery ↔ local /health",
        "After Detect (host PC, server running), header shows Local and the site uses the local API.",
        "Detect fails while the server is running, or header stays Cloud after Detect.",
        "Integration survey question 2",
    ),
    (
        3,
        "Exit local returns the site to Cloud mode",
        "Header ↔ switchToCloudApi",
        "Exit local switches badge to Cloud and cloud features work again.",
        "Exit local does nothing, or the site stays stuck on local/unreachable API.",
        "Integration survey question 3",
    ),
    (
        4,
        "Upload while on Local stores a Local file you can see in the library",
        "Upload ↔ local disk ↔ Library",
        "Upload in Local mode succeeds; library shows the file with a Local badge (or Local filter).",
        "Upload fails, file missing, or still labelled Cloud while Local mode is active.",
        "Integration survey question 4",
    ),
    (
        5,
        "Upload while on Cloud stores a Cloud file you can download",
        "Upload ↔ Cloud storage ↔ Library",
        "Upload in Cloud mode succeeds; badge Cloud; download works.",
        "Upload or download fails, or file labelled Local while on Cloud.",
        "Integration survey question 5",
    ),
    (
        6,
        "LAN address is available to connect other devices",
        "Local setup ↔ /local/share-info",
        "After Detect on the host, a copyable http://LAN-IP:8080 address is shown for other devices.",
        "No LAN address, only 127.0.0.1 for phones, or wrong/unusable IP.",
        "Integration survey question 6",
    ),
    (
        7,
        "Register trusted network on Local Wi‑Fi (first user = admin)",
        "Network settings ↔ TrustedNetwork",
        "First registrant becomes admin; network shows as registered/trusted on that Wi‑Fi.",
        "Cannot register from Local, or cloud-only registration wrongly succeeds for home Wi‑Fi.",
        "Integration survey question 7",
    ),
    (
        8,
        "Share link works for Shared files and not for Private",
        "Share ↔ access mode ↔ /shared",
        "Shared file: create link → open in another tab works; Private: share disabled or blocked.",
        "Private can be shared, or Shared link never opens.",
        "Integration survey question 8",
    ),
    (
        9,
        "Delete keeps a deletion log with who uploaded / deleted",
        "Library ↔ soft delete",
        "File leaves the active list; deletion log shows uploader and deleter.",
        "No deletion log, or names missing.",
        "Integration survey question 9",
    ),
    (
        10,
        "Dashboard Local Network setup button opens the setup page",
        "Dashboard ↔ LocalNetworkSetup",
        "Dashboard button opens Local Network setup successfully.",
        "Button missing or goes to the wrong page.",
        "Integration survey question 10",
    ),
]

for row in integ:
    add_case(*row)

# -------- E2E --------
add_heading_like("C. End-to-End Tests (top of pyramid)", bold=True, size=12)
add_normal(
    "Focus: Complete user workflows; multi-scenario; real-world usage simulation."
)
doc.add_paragraph()

e2e = [
    (
        1,
        "Full cloud journey: register → activate → login → upload → library → download",
        "E2E / Cloud onboarding",
        "Completes the whole path on Cloud without needing the local server.",
        "Stuck at any step (register, activate, login, upload, library, download).",
        "End-to-end survey question 1",
    ),
    (
        2,
        "Full local journey: Dashboard setup → Detect → upload → Local badge → download",
        "E2E / Local Network Mode",
        "Completes host local path; header Local; new file Local; download works.",
        "Cannot complete Detect/upload/download on local, or mode/badge wrong end-to-end.",
        "End-to-end survey question 2",
    ),
    (
        3,
        "Second device joins via LAN URL (not Detect)",
        "E2E / Multi-device LAN",
        "Other device connects with host LAN URL, logs in, and can use library as expected.",
        "Other device cannot connect, or only works with 127.0.0.1 / Detect.",
        "End-to-end survey question 3",
    ),
    (
        4,
        "Switch Local ↔ Cloud in one sitting without breaking the site",
        "E2E / Mode switching",
        "Can Detect local, use features, Exit local, then use Cloud again successfully.",
        "Mode switch breaks login/library, or cannot return to Cloud.",
        "End-to-end survey question 4",
    ),
    (
        5,
        "Local Only file: works on registered Wi‑Fi, blocked off trusted network",
        "E2E / Local Only access",
        "On registered LAN, Local Only download works; outside trusted network, access is denied.",
        "Local Only always works from anywhere, or never works on LAN.",
        "End-to-end survey question 5",
    ),
    (
        6,
        "Network admin vs other user on same Wi‑Fi",
        "E2E / Network admin",
        "First registrant can change/remove network settings; another account cannot.",
        "Non-admin can change/remove registration, or admin lacks controls.",
        "End-to-end survey question 6",
    ),
    (
        7,
        "Share link end-to-end then revoke",
        "E2E / Sharing",
        "Create Shared link → second browser accesses → revoke → access stops.",
        "Link never works, or still works after revoke.",
        "End-to-end survey question 7",
    ),
    (
        8,
        "Idle logout after ~15 minutes without HomeShare activity",
        "E2E / Session",
        "After about 15 minutes idle on the site, user is logged out and must sign in again.",
        "Never logs out, or logs out while actively using HomeShare.",
        "End-to-end survey question 8",
    ),
    (
        9,
        "Mixed Cloud + Local library: filters and badges support finding files",
        "E2E / Library filters",
        "With files from both modes (if available), filters/badges help find the right ones.",
        "Cannot tell Cloud from Local files in the library.",
        "End-to-end survey question 9",
    ),
    (
        10,
        "Overall HomeShare is usable for a class/home sharing scenario",
        "E2E / Overall usability",
        "Tester can complete their assigned scenarios without getting permanently stuck.",
        "Tester cannot complete core tasks even with the Help page.",
        "End-to-end survey question 10 (+ Overall evidence questions)",
    ),
]

for row in e2e:
    add_case(*row)

# -------- SURVEY --------
doc.add_paragraph()
add_heading_like("Survey Questions (assess the test cases)", bold=True, size=13)
add_normal(
    "Instructions for Google Form: for each item use Multiple choice: Yes / Partially / No. "
    "Optional short comment under each section. Put the Form link in Survey link above."
)
doc.add_paragraph()
add_normal("Scoring rule for tutors/team:", bold=True)
add_normal(
    "Yes = Pass for linked case. Partially = Pass with notes (or Fail if you require full success). No = Fail."
)
doc.add_paragraph()

add_heading_like("Section A — Unit questions", bold=True, size=12)

survey_unit = [
    (1, "When registering or resetting a password, weak passwords were rejected and a valid password was accepted."),
    (2, "I could tell Private / Shared / Local Only apart, and Private did not let me create a share link."),
    (3, "Cloud/Local badges and the All/Cloud/Local filters matched what I expected for my files."),
    (4, "The header badge correctly showed Cloud or Local, and Exit local only appeared in Local mode."),
    (5, "The Help page explained Cloud vs Local clearly, and I could open Local Network setup from Help."),
    (6, "I understood that inactivity on HomeShare can log me out (~15 minutes)."),
]

for num, text in survey_unit:
    p = doc.add_paragraph()
    r = p.add_run(f"{num}. ")
    r.bold = True
    r.font.size = Pt(11)
    p.add_run(f"{text}  [Yes / Partially / No]  (assesses Unit case {num})").font.size = Pt(11)

doc.add_paragraph()
add_heading_like("Section B — Integration questions", bold=True, size=12)

survey_int = [
    (1, "I logged in on Cloud without needing a local server or Detect."),
    (2, "On the host PC, Detect put me into Local mode (header showed Local)."),
    (3, "Exit local returned me to Cloud mode successfully."),
    (4, "While in Local mode, my upload appeared in the library as a Local file."),
    (5, "While in Cloud mode, my upload appeared as Cloud and I could download it."),
    (6, "I could copy a LAN address (http://192.168.x.x:8080) for another device."),
    (7, "I could register the trusted network on Local Wi‑Fi (or saw it already registered)."),
    (8, "Share links worked for Shared files and were not available for Private files."),
    (9, "After deleting a file, a deletion log showed who uploaded and who deleted it."),
    (10, "From the Dashboard, Local Network setup opened the correct setup page."),
]

for num, text in survey_int:
    p = doc.add_paragraph()
    r = p.add_run(f"{num}. ")
    r.bold = True
    r.font.size = Pt(11)
    p.add_run(
        f"{text}  [Yes / Partially / No]  (assesses Integration case {num})"
    ).font.size = Pt(11)

doc.add_paragraph()
add_heading_like("Section C — End-to-end questions", bold=True, size=12)

survey_e2e = [
    (1, "I completed cloud register → activate → login → upload → library → download without a local server."),
    (2, "I completed local setup from Dashboard → Detect → upload → saw Local → downloaded the file."),
    (3, "A second device connected using the LAN URL (not Detect) and could use the library."),
    (4, "I switched between Local and Cloud in one session without the site breaking."),
    (5, "A Local Only file worked on the registered Wi‑Fi and was blocked when not on that trusted network (or I could not complete this — answer No)."),
    (6, "Network admin could change settings; a non-admin on the same Wi‑Fi could not (or N/A → Partially with comment)."),
    (7, "I created a share link, opened it in another browser, then revoked it and access stopped."),
    (8, "After ~15 minutes without using HomeShare, I was logged out (or I did not wait — answer Partially and note)."),
    (9, "When I had (or imagined) both Cloud and Local files, badges/filters helped me tell them apart."),
    (10, "Overall I could finish the assigned HomeShare scenarios without getting permanently stuck."),
]

for num, text in survey_e2e:
    p = doc.add_paragraph()
    r = p.add_run(f"{num}. ")
    r.bold = True
    r.font.size = Pt(11)
    p.add_run(
        f"{text}  [Yes / Partially / No]  (assesses End-to-end case {num})"
    ).font.size = Pt(11)

doc.add_paragraph()
add_heading_like("Section D — Evidence & overall (supports marking)", bold=True, size=12)

extra = [
    (1, "Overall, HomeShare felt usable for class/home file sharing (Likert 1–5). Also pick: A) Cloud only B) Local only C) Both D) Blocked."),
    (2, "What was hardest? A) Login/account B) Local setup/Detect C) Upload/download D) Sharing E) Network registration F) Badges/filters G) Nothing major"),
    (3, "If you answered No to any question above, list the section and question number and what went wrong."),
    (4, "Which 1–3 tasks worked best? List section + question number or a short description."),
    (5, "Device + browser used (e.g. Windows 11 + Chrome; phone + Safari)."),
    (6, "Modes you actually tested: Cloud / Local / Both."),
    (7, "One improvement you would make first (short answer)."),
]

for num, text in extra:
    p = doc.add_paragraph()
    r = p.add_run(f"{num}. ")
    r.bold = True
    r.font.size = Pt(11)
    p.add_run(text).font.size = Pt(11)

doc.add_paragraph()
add_heading_like("How cases map to survey questions", bold=True, size=12)
add_normal("Unit case N ↔ Unit survey question N")
add_normal("Integration case N ↔ Integration survey question N")
add_normal("End-to-end case N ↔ End-to-end survey question N")
add_normal("End-to-end case 10 is also supported by Section D overall evidence questions.")

doc.add_paragraph()
add_normal("Notes for testers:", bold=True)
add_normal("• Answer survey items only for scenarios you actually tried; use Partially + comment if skipped.")
add_normal("• Cloud: live GitHub Pages, no Detect. Local: Start HomeShare.bat on host, Detect on host, LAN URL on other devices.")
add_normal("• Team: transfer Yes/Partial/No from Form responses into Pass/Fail on each matching test case.")

doc.save(out)
print("Wrote", out)
print(
    "Cases:",
    len(unit) + len(integ) + len(e2e),
    "Survey items:",
    len(survey_unit) + len(survey_int) + len(survey_e2e) + len(extra),
)
